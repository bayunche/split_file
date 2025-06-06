import io
import requests
from typing import List, Dict


def split_word_from_url(url: str, chunk_size: int = 500, overlap: int = 50, mode: str = "paragraph") -> List[Dict]:
    """
    从URL读取Word文档并按指定模式切分

    Args:
        url: Word文档URL
        chunk_size: 每个片段的大小（含义取决于mode）
        overlap: 相邻片段的重叠大小
        mode: 切分模式，"paragraph"按段落数，"character"按字符数，"line"按换行符

    Returns:
        切分后的片段列表，每个片段包含元信息和文本内容
    """
    try:
        # 下载文档
        response = requests.get(url)
        response.raise_for_status()

        # 尝试导入并使用 python-docx
        try:
            from docx import Document
        except ImportError:
            return [{"id": 0, "text": "依赖错误: 无法导入 python-docx 库。请确保已安装 python-docx==0.8.11", "error": True}]

        # 解析Word文档
        try:
            doc = Document(io.BytesIO(response.content))
        except Exception as e:
            return [{"id": 0, "text": f"文档解析失败: {str(e)}。请确保文件是有效的Word文档(.docx格式)。", "error": True}]

        # 根据模式选择切分方法
        if mode.lower() == "paragraph":
            # 按段落切分
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            if not paragraphs:
                return [{"id": 0, "text": "文档无有效内容", "start": 0, "end": 0}]

            chunks = []
            total_paragraphs = len(paragraphs)

            # 计算需要切分的部分数量
            parts = max(1, (total_paragraphs - overlap) // (chunk_size - overlap))

            for i in range(parts):
                start_idx = i * (chunk_size - overlap)
                end_idx = min(start_idx + chunk_size, total_paragraphs)

                chunks.append({
                    "id": i + 1,
                    "text": "\n".join(paragraphs[start_idx:end_idx]),
                    "start": start_idx + 1,
                    "end": end_idx,
                    "total": total_paragraphs,
                    "mode": "paragraph",
                    "size": end_idx - start_idx
                })

            # 处理剩余段落
            if end_idx < total_paragraphs:
                chunks.append({
                    "id": len(chunks) + 1,
                    "text": "\n".join(paragraphs[end_idx:]),
                    "start": end_idx + 1,
                    "end": total_paragraphs,
                    "total": total_paragraphs,
                    "mode": "paragraph",
                    "size": total_paragraphs - end_idx
                })

        elif mode.lower() == "character":
            # 按字符数切分
            full_text = "\n".join([p.text for p in doc.paragraphs])
            if not full_text.strip():
                return [{"id": 0, "text": "文档无有效内容", "start": 0, "end": 0}]

            chunks = []
            total_chars = len(full_text)

            # 计算需要切分的部分数量
            parts = max(1, (total_chars - overlap) // (chunk_size - overlap))

            for i in range(parts):
                start_idx = i * (chunk_size - overlap)
                end_idx = min(start_idx + chunk_size, total_chars)

                # 尝试在段落边界处切分，避免截断句子
                if end_idx < total_chars:
                    # 寻找最近的段落结束符
                    last_newline = full_text[:end_idx].rfind('\n')
                    if last_newline > start_idx + chunk_size * 0.8:  # 确保不切得太短
                        end_idx = last_newline

                chunks.append({
                    "id": i + 1,
                    "text": full_text[start_idx:end_idx],
                    "start": start_idx,
                    "end": end_idx,
                    "total": total_chars,
                    "mode": "character",
                    "size": end_idx - start_idx
                })

            # 处理剩余字符
            if end_idx < total_chars:
                chunks.append({
                    "id": len(chunks) + 1,
                    "text": full_text[end_idx:],
                    "start": end_idx,
                    "end": total_chars,
                    "total": total_chars,
                    "mode": "character",
                    "size": total_chars - end_idx
                })

        elif mode.lower() == "line":
            # 按换行符切分并结合长度限制
            full_text = "\n".join([p.text for p in doc.paragraphs])
            if not full_text.strip():
                return [{"id": 0, "text": "文档无有效内容", "start": 0, "end": 0}]

            # 按换行符分割成多个块
            lines = full_text.split('\n')
            non_empty_lines = [line for line in lines if line.strip()]

            if not non_empty_lines:
                return [{"id": 0, "text": "文档无有效内容", "start": 0, "end": 0}]

            chunks = []
            current_chunk = []
            current_length = 0
            total_lines = len(non_empty_lines)

            for i, line in enumerate(non_empty_lines):
                line_length = len(line)

                # 如果添加当前行后会超过 chunk_size，且当前块已有内容，则创建新片段
                if current_length + line_length > chunk_size and current_chunk:
                    chunks.append({
                        "id": len(chunks) + 1,
                        "text": "\n".join(current_chunk),
                        "start": i - len(current_chunk),
                        "end": i - 1,
                        "total": total_lines,
                        "mode": "line",
                        "size": current_length
                    })

                    # 添加重叠内容
                    overlap_lines = max(1, len(current_chunk) * overlap // chunk_size)
                    if overlap_lines > 0 and len(chunks) > 0:
                        current_chunk = current_chunk[-overlap_lines:]
                        current_length = sum(len(l) for l in current_chunk)
                    else:
                        current_chunk = []
                        current_length = 0

                # 添加当前行到当前片段
                current_chunk.append(line)
                current_length += line_length

            # 添加最后一个片段
            if current_chunk:
                chunks.append({
                    "id": len(chunks) + 1,
                    "text": "\n".join(current_chunk),
                    "start": total_lines - len(current_chunk),
                    "end": total_lines - 1,
                    "total": total_lines,
                    "mode": "line",
                    "size": current_length
                })

        else:
            return [{"id": 0, "text": f"不支持的切分模式: {mode}。请使用 'paragraph', 'character' 或 'line'。", "error": True}]

        return chunks

    except Exception as e:
        return [{"id": 0, "text": f"处理失败: {str(e)}", "error": True}]


# 平台要求的入口函数 - 兼容单参数和双参数调用方式
def handler(*args):
    """
    Coze平台要求的入口函数

    Args:
        可能是单参数 (event) 或双参数 (event, context)

    Returns:
        处理结果
    """
    try:
        # 打印原始参数（用于调试）
        print(f"接收到的原始参数: {args}")

        # 提取 event 参数（兼容单参数和双参数调用方式）
        if not args:
            return [{"id": 0, "text": "未提供参数", "error": True}]

        event = args[0]

        # 处理特殊的 Args/CustomNamespace 包装结构
        url = None
        chunk_size = 500
        overlap = 50
        mode = "paragraph"

        # 检查是否为 CustomNamespace 对象
        if hasattr(event, 'input') and hasattr(event.input, 'url'):
            url = event.input.url
            chunk_size = getattr(event.input, 'chunkSize', 500)
            overlap = getattr(event.input, 'overlap', 50)
            mode = getattr(event.input, 'mode', "paragraph")

            # 将字符串类型的数字转换为整数
            if isinstance(chunk_size, str):
                try:
                    chunk_size = int(chunk_size)
                except ValueError:
                    return [{"id": 0, "text": f"无效的 chunkSize 参数: {chunk_size}，必须是整数", "error": True}]

            if isinstance(overlap, str):
                try:
                    overlap = int(overlap)
                except ValueError:
                    return [{"id": 0, "text": f"无效的 overlap 参数: {overlap}，必须是整数", "error": True}]
        else:
            # 常规参数提取方式
            try:
                url = event.get("url")
                chunk_size = event.get("chunk_size", 500)
                overlap = event.get("overlap", 50)
                mode = event.get("mode", "paragraph")
            except AttributeError:
                url = getattr(event, "url", None)
                chunk_size = getattr(event, "chunk_size", 500)
                overlap = getattr(event, "overlap", 50)
                mode = getattr(event, "mode", "paragraph")

        # 打印解析后的参数（用于调试）
        print(f"解析后的参数: url={url}, chunk_size={chunk_size}, overlap={overlap}, mode={mode}")

        # 参数验证
        if not url:
            return [{"id": 0, "text": f"缺少必要参数: url。原始参数: {args}", "error": True}]

        # 调用实际处理函数
        return split_word_from_url(url, chunk_size, overlap, mode)

    except Exception as e:
        return [{"id": 0, "text": f"处理过程中发生意外错误: {str(e)}", "error": True}]